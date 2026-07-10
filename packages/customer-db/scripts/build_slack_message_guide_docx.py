#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SLACK_MESSAGE_GUIDE.md"
OUTPUT = ROOT / "docs" / "Slack_Message_Guide_tf_cross_team_sales.docx"


FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 101, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None, name=FONT) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text_with_code(paragraph, text: str) -> None:
    parts = text.split("`")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2:
            set_font(run, size=10, color=DARK_BLUE, name="Consolas")
        else:
            set_font(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="DADCE0", size="3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    text = "\n".join(lines).rstrip()
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        set_font(run, size=9.5, name="Consolas")
    doc.add_paragraph()


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(1.25), Inches(2.65), Inches(2.60)] if len(rows[0]) == 3 else None
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = widths[c_idx]
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_text_with_code(para, value.strip())
            for run in para.runs:
                set_font(run, size=9.5, bold=(r_idx == 0), color=DARK_BLUE if r_idx == 0 else None)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(set(c.replace(" ", "")) <= {"-", ":"} for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.line_spacing = 1.1
    run = title.add_run("Slack 메시지 작성 가이드")
    set_font(run, size=24, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("#tf_cross_team_sales · 고객 DB 자동 정리를 위한 공유용 안내")
    set_font(run, size=11, color=MUTED)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run("강제 입력 양식이 아니라, Slack 메시지가 GLM을 거쳐 고객 DB에 더 잘 정리되도록 돕는 가벼운 작성 제안입니다.")
    set_font(run, size=11, bold=True)


def build() -> None:
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_markdown_table(lines, i)
            add_markdown_table(doc, rows)
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_code(p, stripped[2:])
            i += 1
            continue

        if len(stripped) > 3 and stripped[:2].isdigit() and stripped[2] == ".":
            p = doc.add_paragraph(style="List Number")
            add_text_with_code(p, stripped[3:].strip())
            i += 1
            continue

        p = doc.add_paragraph()
        add_text_with_code(p, stripped.replace("  ", " "))
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("#tf_cross_team_sales 메시지 작성 가이드")
    set_font(run, size=8.5, color=MUTED)

    doc.core_properties.title = "Slack 메시지 작성 가이드"
    doc.core_properties.subject = "#tf_cross_team_sales 고객 DB 자동 정리 가이드"
    doc.core_properties.author = "RTM"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

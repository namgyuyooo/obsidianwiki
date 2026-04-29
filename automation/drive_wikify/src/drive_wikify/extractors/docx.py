from __future__ import annotations

from pathlib import Path

from ..models import ExtractedContent


def extract_docx(path: Path) -> ExtractedContent:
    import docx

    document = docx.Document(str(path))
    paragraphs = []
    headings = []
    tables = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if paragraph.style and "Heading" in paragraph.style.name:
            headings.append(text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
            flat = " | ".join(" / ".join(cell for cell in row if cell) for row in rows[:5])
            if flat:
                paragraphs.append(flat)

    return ExtractedContent(
        extractor_name="python_docx",
        text="\n".join(paragraphs),
        headings=headings,
        tables=tables,
    )
